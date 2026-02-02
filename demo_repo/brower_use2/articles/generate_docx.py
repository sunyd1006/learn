from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import requests
from io import BytesIO

def set_heading_styles(doc):
    """设置标题样式"""
    # 设置一级标题样式
    doc.styles['Heading 1'].font.size = Pt(16)
    doc.styles['Heading 1'].font.bold = True
    doc.styles['Heading 1'].paragraph_format.space_before = Pt(12)
    doc.styles['Heading 1'].paragraph_format.space_after = Pt(6)
    
    # 设置二级标题样式
    doc.styles['Heading 2'].font.size = Pt(14)
    doc.styles['Heading 2'].font.bold = True
    doc.styles['Heading 2'].paragraph_format.space_before = Pt(10)
    doc.styles['Heading 2'].paragraph_format.space_after = Pt(4)

def add_image_from_url(doc, url, width=Inches(6)):
    """从URL添加图片到文档"""
    try:
        response = requests.get(url, timeout=10)
        response.raise_for_status()
        image_stream = BytesIO(response.content)
        doc.add_picture(image_stream, width=width)
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return True
    except Exception as e:
        print(f"添加图片失败: {e}")
        return False

def main():
    # 创建文档
    doc = Document()
    
    # 设置标题样式
    set_heading_styles(doc)
    
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)

    # 文档标题
    title = doc.add_heading('人工智能技术应用与发展报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.style.font.size = Pt(20)
    title.style.font.bold = True

    doc.add_paragraph()  # 空行

    # 摘要
    abstract_heading = doc.add_paragraph('摘要')
    abstract_heading.bold = True
    abstract = doc.add_paragraph(
        "本报告全面介绍了人工智能技术的发展现状、核心应用领域以及未来发展趋势。"
        "通过对大语言模型、计算机视觉、自然语言处理等关键技术的分析，探讨了人工智能"
        "在各行业的落地场景和实际价值，同时也讨论了技术发展过程中面临的挑战和伦理问题。"
    )
    abstract.paragraph_format.first_line_indent = Inches(0.3)

    doc.add_page_break()

    # 第一章：人工智能发展概述
    heading1 = doc.add_heading('1. 人工智能发展概述', level=1)

    # 1.1 发展历程
    heading2 = doc.add_heading('1.1 发展历程', level=2)

    para1 = doc.add_paragraph(
        "人工智能的发展可以追溯到20世纪50年代。1956年的达特茅斯会议正式提出了"
        "'人工智能'这一概念，标志着人工智能学科的诞生。在随后的几十年里，人工智能"
        "经历了多次兴衰起伏，从早期的逻辑推理、专家系统，到机器学习时代，再到当前"
        "的深度学习革命阶段。"
    )
    para1.paragraph_format.first_line_indent = Inches(0.3)

    para2 = doc.add_paragraph(
        "2012年，AlexNet在ImageNet图像识别竞赛中取得突破性成绩，深度学习技术开始"
        "快速发展。2022年底ChatGPT的发布，更是将大语言模型技术推向了公众视野，"
        "人工智能技术开始进入大规模应用阶段。"
    )
    para2.paragraph_format.first_line_indent = Inches(0.3)

    # 添加AI发展时间线图片
    add_image_from_url(doc, "https://picsum.photos/id/0/800/400")
    doc.add_paragraph("图1-1 人工智能发展历程时间线").alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 1.2 核心技术体系
    heading2 = doc.add_heading('1.2 核心技术体系', level=2)

    para3 = doc.add_paragraph(
        "当前人工智能技术体系主要包括以下几个核心领域："
    )
    para3.paragraph_format.first_line_indent = Inches(0.3)

    doc.add_paragraph("• 大语言模型：具备理解和生成人类语言的能力，是当前最受关注的技术方向", style='List Bullet')
    doc.add_paragraph("• 计算机视觉：让机器具备'看'的能力，可应用于图像识别、视频分析等场景", style='List Bullet')
    doc.add_paragraph("• 语音技术：包括语音识别和语音合成，实现人与机器的语音交互", style='List Bullet')
    doc.add_paragraph("• 多模态技术：融合文本、图像、音频、视频等多种信息的处理能力", style='List Bullet')
    doc.add_paragraph("• 强化学习：通过与环境交互不断优化决策能力，在机器人、游戏等领域应用广泛", style='List Bullet')

    doc.add_page_break()

    # 第二章：人工智能典型应用场景
    heading1 = doc.add_heading('2. 人工智能典型应用场景', level=1)

    # 2.1 企业服务领域
    heading2 = doc.add_heading('2.1 企业服务领域', level=2)

    para4 = doc.add_paragraph(
        "在企业服务领域，人工智能技术被广泛应用于智能客服、人力资源管理、财务自动化、"
        "供应链优化等场景。智能客服系统可以自动回答70%以上的常见问题，大幅降低企业"
        "的客服成本，同时提升响应速度和用户体验。"
    )
    para4.paragraph_format.first_line_indent = Inches(0.3)

    # 2.2 医疗健康领域
    heading2 = doc.add_heading('2.2 医疗健康领域', level=2)

    para5 = doc.add_paragraph(
        "在医疗健康领域，人工智能在医学影像诊断、药物研发、辅助诊疗、健康管理等方面"
        "展现出巨大价值。AI辅助诊断系统可以帮助医生更快速、更准确地识别CT、MRI等"
        "医学影像中的病灶，提高诊断效率和准确率。"
    )
    para5.paragraph_format.first_line_indent = Inches(0.3)

    # 添加医疗AI应用图片
    add_image_from_url(doc, "https://picsum.photos/id/1/800/500")
    doc.add_paragraph("图2-1 AI辅助医学影像诊断应用示例").alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 2.3 智能制造领域
    heading2 = doc.add_heading('2.3 智能制造领域', level=2)

    para6 = doc.add_paragraph(
        "在智能制造领域，人工智能技术推动工业生产向智能化、自动化方向发展。"
        " predictive maintenance系统可以通过分析设备传感器数据，提前预测设备故障，"
        "减少非计划停机时间。质量检测AI系统可以实时识别生产线上的缺陷产品，"
        "提高产品质量和生产效率。"
    )
    para6.paragraph_format.first_line_indent = Inches(0.3)

    doc.add_page_break()

    # 第三章：人工智能发展趋势与挑战
    heading1 = doc.add_heading('3. 人工智能发展趋势与挑战', level=1)

    # 3.1 未来发展趋势
    heading2 = doc.add_heading('3.1 未来发展趋势', level=2)

    para7 = doc.add_paragraph(
        "未来人工智能技术将呈现以下发展趋势：一是模型向多模态、通用化方向发展，"
        "具备更强的通用智能能力；二是边缘AI技术快速发展，AI推理和计算将更多在"
        "终端设备上运行；三是AI与行业场景深度融合，垂直领域的专用模型将大量涌现；"
        "四是Agent技术逐渐成熟，自主智能体将能够完成更复杂的任务。"
    )
    para7.paragraph_format.first_line_indent = Inches(0.3)

    # 添加AI未来趋势图片
    add_image_from_url(doc, "https://picsum.photos/id/3/800/450")
    doc.add_paragraph("图3-1 人工智能未来技术发展路线").alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 3.2 面临的挑战
    heading2 = doc.add_heading('3.2 面临的挑战', level=2)

    para8 = doc.add_paragraph(
        "在人工智能快速发展的同时，也面临着诸多挑战：一是伦理和安全问题，如何确保"
        "AI系统的公平性、透明性和安全性；二是数据隐私问题，AI训练需要大量数据，"
        "如何在数据利用和隐私保护之间取得平衡；三是就业结构调整，AI的广泛应用"
        "可能会对部分职业造成冲击；四是技术治理问题，需要建立健全AI相关的法律法规"
        "和监管体系。"
    )
    para8.paragraph_format.first_line_indent = Inches(0.3)

    # 结论
    doc.add_page_break()
    conclusion_heading = doc.add_paragraph('结论')
    conclusion_heading.runs[0].bold = True
    conclusion_heading.runs[0].font.size = Pt(14)
    
    conclusion = doc.add_paragraph(
        "人工智能作为新一轮科技革命和产业变革的核心驱动力，正在深刻改变人类的生产"
        "和生活方式。在把握技术发展机遇的同时，我们也需要积极应对各种挑战，推动"
        "人工智能技术朝着更加安全、可控、普惠的方向发展，让技术进步更好地造福人类社会。"
    )
    conclusion.paragraph_format.first_line_indent = Inches(0.3)

    # 保存文档
    output_path = '/Users/bytedance/codespace/learn/demo_repo/brower_use2/articles/人工智能技术发展报告.docx'
    doc.save(output_path)
    print(f"文档已成功生成并保存到: {output_path}")

if __name__ == "__main__":
    main()
